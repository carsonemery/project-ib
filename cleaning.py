import pandas as pd
import json
import re

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX export will not be available.")


def clean_text(text):
    """Clean whitespace and normalize text formatting."""
    if not text:
        return ""
    # Remove excessive whitespace while preserving intentional line breaks
    text = re.sub(r'\n+', '\n', text)  # Multiple newlines to single
    text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
    return text.strip()


def load_technical_flashcards(filename="technical_flashcards_complete_v2.json"):
    """Load technical flashcards and parse tags using smart parsing."""
    with open(filename, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    rows = []
    for obj in data:
        tags = obj.get("tags", [])
        
        # Smart parsing approach
        reported_in = ""
        question_type = ""
        
        # Find reported_in (should be first match from REPORTED_IN_KEYWORDS)
        for tag in tags:
            if tag in REPORTED_IN_KEYWORDS:
                reported_in = tag
                break
        
        # Find question_type (first match from TYPE_CATEGORIES)  
        for tag in tags:
            if tag in TYPE_CATEGORIES:
                question_type = tag
                break
        
        row = {
            "Question Number": obj.get("questionNumber", ""),
            "Question": clean_text(obj.get("question", "")),
            "Reported In": reported_in,
            "Type": question_type,
            "Answer": clean_text(obj.get("answer", "")),
            "Source": "Technical"
        }
        rows.append(row)
    
    return pd.DataFrame(rows)

# Known categories from tag analysis
REPORTED_IN_KEYWORDS = {
    "IB Vine Select",
    "Reported in Allen & Company Interview",
    "Reported in Ardea Partners Interview", 
    "Reported in Ares Management Interview",
    "Reported in BMO Interview",
    "Reported in Bank of America Interview",
    "Reported in Barclays Interview",
    "Reported in Brookfield Asset Management Interview",
    "Reported in CIBC Capital Markets Interview",
    "Reported in Centerview Partners Interview",
    "Reported in Citi Interview",
    "Reported in Deutsche Bank Interview",
    "Reported in Evercore Interview",
    "Reported in FT Partners Interview",
    "Reported in Goldman Sachs Interview",
    "Reported in Greenhill & Co. Interview",
    "Reported in Guggenheim Partners Interview",
    "Reported in Houlihan Lokey Interview",
    "Reported in J.P. Morgan Interview",
    "Reported in Jefferies Interview",
    "Reported in Lazard Interview",
    "Reported in Leerink Partners Interview",
    "Reported in LionTree Interview",
    "Reported in M. Klein & Co. Interview",
    "Reported in Moelis & Co. Interview",
    "Reported in Morgan Stanley Interview",
    "Reported in Nomura Greentech Interview",
    "Reported in PJT Partners Interview",
    "Reported in Perella Weinberg Interview",
    "Reported in Piper Sandler Interview",
    "Reported in Qatalyst Partners Interview",
    "Reported in RBC Interview",
    "Reported in Raine Group Interview",
    "Reported in Rothschild Interview",
    "Reported in UBS Interview",
    "Reported in Union Square Advisors Interview",
    "Reported in Wells Fargo Interview",
    "Reported in William Blair Interview"
}

# Combined type categories from positions 1 and 2
TYPE_CATEGORIES = {
    "Accounting",
    "DCF", 
    "Enterprise & Equity Value",
    "Finance Brain Teasers",
    "General Brain Teasers",
    "Leveraged Buyouts (LBOs)",
    "Market Knowledge",
    "Merger Model (M&A)",
    "Restructuring / Distressed M&A",
    "Restructuring group",
    "Technology Banking",
    "Valuation"
}

def load_all_flashcards():
    """Load technical flashcards."""
    return load_technical_flashcards()


def format_answer_with_bullets(answer_text):
    """Convert answer text to proper bullet points for DOCX."""
    if not answer_text:
        return []
    
    lines = answer_text.split('\n')
    formatted_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if line already has bullet
        if line.startswith('•'):
            formatted_lines.append(('bullet', line[1:].strip()))
        elif line.startswith('-') and len(line) > 1 and line[1] == ' ':
            formatted_lines.append(('bullet', line[2:].strip()))
        else:
            formatted_lines.append(('normal', line))
    
    return formatted_lines


def export_to_docx(df, filename="flashcards.docx", filter_by=None):
    """
    Export DataFrame to DOCX with proper formatting.
    
    Args:
        df: DataFrame with flashcard data
        filename: Output DOCX filename
        filter_by: Dict to filter data (e.g., {"Type": "Accounting"})
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
    
    # Filter data if specified
    filtered_df = df.copy()
    if filter_by:
        for column, value in filter_by.items():
            if column in filtered_df.columns:
                filtered_df = filtered_df[filtered_df[column] == value]
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('Investment Banking Flashcards', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    if filter_by:
        subtitle = doc.add_heading(f"Filtered by: {', '.join([f'{k}={v}' for k, v in filter_by.items()])}", level=2)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add summary
    doc.add_paragraph(f"Total Questions: {len(filtered_df)}")
    doc.add_paragraph("")
    
    # Add each flashcard
    for idx, row in filtered_df.iterrows():
        # Question section
        question_para = doc.add_paragraph()
        question_run = question_para.add_run("Question: ")
        question_run.bold = True
        question_para.add_run(str(row.get('Question', '')))
        
        # Reported In section (if applicable)
        if row.get('Reported In'):
            reported_para = doc.add_paragraph()
            reported_run = reported_para.add_run("Reported In: ")
            reported_run.bold = True
            reported_para.add_run(str(row.get('Reported In', '')))
        
        # Type (if applicable)
        if row.get('Type'):
            type_para = doc.add_paragraph()
            type_run = type_para.add_run("Type: ")
            type_run.bold = True
            type_para.add_run(str(row.get('Type', '')))
        
        # Answer section
        answer_para = doc.add_paragraph()
        answer_run = answer_para.add_run("Answer:")
        answer_run.bold = True
        
        # Format answer with bullets
        answer_lines = format_answer_with_bullets(str(row.get('Answer', '')))
        for line_type, line_text in answer_lines:
            if line_type == 'bullet':
                bullet_para = doc.add_paragraph(line_text, style='List Bullet')
            else:
                doc.add_paragraph(line_text)
        
        # Add separator
        doc.add_paragraph("─" * 50)
        doc.add_paragraph("")
    
    # Save document
    doc.save(filename)
    print(f"DOCX exported successfully to: {filename}")


def export_by_categories(df, output_dir="./"):
    """Export separate DOCX files for each major category."""
    if not DOCX_AVAILABLE:
        print("DOCX export not available. Install python-docx first.")
        return
    
    # Get unique types
    unique_types = df['Type'].dropna().unique()
    
    for category in unique_types:
        if category:  # Skip empty categories
            safe_filename = re.sub(r'[^\w\-_\.]', '_', category)
            filename = f"{output_dir}flashcards_{safe_filename}.docx"
            export_to_docx(df, filename, filter_by={"Type": category})
            print(f"Exported {category} to {filename}")


def export_by_reported_in(df, output_dir="./"):
    """Export separate DOCX files for each firm/source."""
    if not DOCX_AVAILABLE:
        print("DOCX export not available. Install python-docx first.")
        return
    
    unique_reported = df['Reported In'].dropna().unique()
    
    for reported in unique_reported:
        if reported:  # Skip empty reported
            safe_filename = re.sub(r'[^\w\-_\.]', '_', reported)
            filename = f"{output_dir}flashcards_{safe_filename}.docx"
            export_to_docx(df, filename, filter_by={"Reported In": reported})
            print(f"Exported {reported} questions to {filename}")


# Run
if __name__ == "__main__":
    print("Loading technical flashcards...")
    combined = load_all_flashcards()
    print(f"Total technical questions: {len(combined)}")
    
    print("\nExporting by type/category...")
    export_by_categories(combined)
    










